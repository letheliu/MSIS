from abc import ABC, abstractmethod
from pathlib import Path
import json
from typing import Optional, Any
from docx import Document
from docx.opc.exceptions import PackageNotFoundError
import pdfplumber


class DocumentParser(ABC):
    """文档解析器抽象基类"""

    @abstractmethod
    def parse(self, file_path: Path) -> str:
        """解析文档并返回格式化文本"""
        pass


class TextParser(DocumentParser):
    """文本文件解析器 (txt, md, json)"""

    def parse(self, file_path: Path) -> str:
        try:
            suffix = file_path.suffix.lower()

            if suffix == '.json':
                with open(file_path, 'r', encoding='utf-8') as f:
                    data = json.load(f)
                    return self._extract_json_strings(data)

            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
        except FileNotFoundError:
            raise FileNotFoundError(f"File not found: {file_path}")
        except json.JSONDecodeError as e:
            raise ValueError(f"Invalid JSON file: {file_path}. Error: {e}")
        except UnicodeDecodeError as e:
            raise ValueError(f"Encoding error reading file {file_path}. Error: {e}")
        except IOError as e:
            raise IOError(f"Error reading file {file_path}: {e}")

    def _extract_json_strings(self, data) -> str:
        """递归提取 JSON 中的所有字符串值"""
        if isinstance(data, str):
            return data
        elif isinstance(data, dict):
            return ' '.join(self._extract_json_strings(v) for v in data.values())
        elif isinstance(data, list):
            return ' '.join(self._extract_json_strings(item) for item in data)
        return ''


class DocxParser(DocumentParser):
    """Word 文档解析器"""

    def parse(self, file_path: Path) -> str:
        try:
            # Document() expects a string path, not Path object
            doc = Document(str(file_path))
            content = []

            for para in doc.paragraphs:
                text = para.text.strip()
                if not text:
                    continue

                # 检查样式是否为标题 - para.style can be None
                style_name: Optional[str] = para.style.name if para.style else None
                if style_name and style_name.startswith('Heading'):
                    level = style_name.replace('Heading ', '')
                    try:
                        hashes = '#' * min(int(level) + 1, 6)
                        content.append(f"{hashes} {text}")
                    except ValueError:
                        # If level is not a valid integer, treat as regular text
                        content.append(text)
                else:
                    content.append(text)

            # 处理表格
            for table in doc.tables:
                content.append("")
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells]
                    content.append(" | ".join(cells))
                content.append("")

            return "\n".join(content)
        except (FileNotFoundError, PackageNotFoundError):
            raise FileNotFoundError(f"File not found: {file_path}")
        except IOError as e:
            raise IOError(f"Error reading DOCX file {file_path}: {e}")
        except Exception as e:
            raise RuntimeError(f"Error parsing DOCX file {file_path}: {e}")


class PDFParser(DocumentParser):
    """PDF 文档解析器"""

    def parse(self, file_path: Path) -> str:
        try:
            content = []

            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    text = page.extract_text()
                    if text:
                        content.append(text)

                    # 尝试提取表格
                    tables = page.extract_tables()
                    if tables:
                        for table in tables:
                            content.append("")
                            for row in table:
                                cells = [str(cell).strip() if cell else "" for cell in row]
                                content.append(" | ".join(cells))
                            content.append("")

            return "\n".join(content)
        except FileNotFoundError:
            raise FileNotFoundError(f"File not found: {file_path}")
        except IOError as e:
            raise IOError(f"Error reading PDF file {file_path}: {e}")
        except Exception as e:
            raise RuntimeError(f"Error parsing PDF file {file_path}: {e}")
